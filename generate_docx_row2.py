# -*- coding: utf-8 -*-
"""
Codiste LinkedIn Article — DOCX Generator
Row 2: AI in RegTech - Zero-Trust Compliance C-Suite Blueprint
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUT_DIR = '/home/user/linkedin_post_creation/output/zero-trust-compliance-ai-regtech/'
os.makedirs(OUT_DIR, exist_ok=True)

TITLE = "Zero-Trust Compliance: The C-Suite Playbook for 2026"
HASHTAGS_ARTICLE = "#ZeroTrust #RegTech #AICompliance #FinancialServices #CyberSecurity"


def unicode_bold(text):
    result = ''
    for c in text:
        if 'A' <= c <= 'Z':
            result += chr(0x1D5D4 + ord(c) - ord('A'))
        elif 'a' <= c <= 'z':
            result += chr(0x1D5EE + ord(c) - ord('a'))
        else:
            result += c
    return result


CAPTION_LINE1 = unicode_bold(
    "Your security team is solving a breach problem. Your compliance team is solving a reporting "
    "problem. Neither is solving the architecture problem they share."
)
CAPTION_LINE2 = (
    "The C-suite blueprint for merging zero-trust and RegTech into one AI-driven compliance layer "
    "→ [article link]"
)
CAPTION_HASHTAGS = "#ZeroTrustSecurity #AIRegTech #FinancialCompliance #CISOInsights #RegulatoryTechnology"


def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D7D7D7')
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_divider(doc, label):
    div1 = doc.add_paragraph()
    div1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = div1.add_run("=" * 50)
    set_font(r1, size=10, color=(100, 100, 100))
    lbl_p = doc.add_paragraph()
    lbl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = lbl_p.add_run(label)
    set_font(rl, size=11, bold=True, color=(60, 60, 60))
    div2 = doc.add_paragraph()
    div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = div2.add_run("=" * 50)
    set_font(r2, size=10, color=(100, 100, 100))
    doc.add_paragraph()


def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=14, bold=True, color=(1, 1, 60))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=11)
    p.paragraph_format.space_after = Pt(8)


def bullet(doc, lead, rest=""):
    p = doc.add_paragraph(style='List Bullet')
    if lead:
        rb = p.add_run(lead + (":" if rest else ""))
        set_font(rb, size=11, bold=True)
    if rest:
        rr = p.add_run(rest)
        set_font(rr, size=11)
    p.paragraph_format.space_after = Pt(5)


def build_docx():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    section_divider(doc, "\U0001f537 LINKEDIN ARTICLE — CODISTE")

    lbl = doc.add_paragraph()
    r = lbl.add_run("TITLE:")
    set_font(r, size=10, bold=True, color=(120, 120, 120))

    tp = doc.add_paragraph()
    tr = tp.add_run(TITLE)
    set_font(tr, size=20, bold=True, color=(1, 1, 1))
    doc.add_paragraph()

    lbl2 = doc.add_paragraph()
    r2 = lbl2.add_run("KEYWORDS TO USE AS HASHTAGS:")
    set_font(r2, size=10, bold=True, color=(120, 120, 120))

    hp = doc.add_paragraph()
    hr = hp.add_run(HASHTAGS_ARTICLE)
    set_font(hr, size=11, color=(0, 100, 200))
    doc.add_paragraph()

    section_divider(doc, "ARTICLE BODY")

    # HOOK
    hook_p = doc.add_paragraph()
    hook_r = hook_p.add_run(
        "84% of organisations experienced identity-related breaches in 2025. Most weren’t "
        "technology failures — they were architecture failures. Systems that verified once "
        "at login and trusted everything after.\n\n"
        "For regulated industries, a single breach isn’t just a security event. It’s a "
        "regulatory one. Zero-trust compliance fixes the architecture. AI makes it scale without "
        "adding headcount."
    )
    set_font(hook_r, size=12, bold=True)
    hook_p.paragraph_format.space_after = Pt(12)

    # SECTION 1: Why AI in RegTech Is Now a Zero-Trust Imperative
    heading(doc, "Why AI in RegTech Is Now a Zero-Trust Imperative")

    para(doc,
        "Traditional RegTech focused on automated reporting and static dashboards. Security teams "
        "handled zero-trust in separate silos. Today’s reality demands a different approach."
    )
    para(doc,
        "AI-enhanced zero-trust controls platforms are merging these two worlds — using AI to "
        "continuously re-evaluate access, detect risky behaviour in real time, and trigger compliance "
        "alerts when patterns match risk thresholds."
    )
    para(doc,
        "For the C-suite, zero-trust compliance is no longer a network security story. It’s a "
        "governance, risk, and compliance imperative. If your security and compliance teams still "
        "operate in separate silos, your RegTech investment is solving only half the problem. "
        "Zero-trust and GRC need a shared data layer, not just a shared meeting."
    )

    # SECTION 2: The C-Suite Blueprint
    heading(doc, "The C-Suite Blueprint: Three Shifts That Change Everything")

    para(doc,
        "Aligning AI in RegTech with zero-trust compliance requires three executive-level shifts."
    )

    para(doc, "Re-frame zero-trust as a compliance enabler, not a security cost:")
    bullet(doc, "Every transaction, session, and user",
        " is verified continuously — not just at login")
    bullet(doc, "All access decisions",
        " are logged, creating the audit trail regulators require on demand")
    bullet(doc, "Compromised credential blast radius",
        " is reduced, directly lowering regulatory risk exposure")

    para(doc, "Embed AI-driven controls into core workflows:")
    bullet(doc, "Contextual access",
        ": AI dynamically adjusts permissions based on role, location, device, and real-time behaviour")
    bullet(doc, "Real-time monitoring",
        ": AI models flag suspicious patterns before they become incidents — not in a quarterly report")
    bullet(doc, "Human-risk scoring dashboards",
        ": identity, activity, and security telemetry combined into a prioritised queue")

    para(doc, "Standardise governance and vendor selection:")
    bullet(doc, "Does the platform support AI-driven identity governance",
        " aligned with zero-trust compliance mandates?")
    bullet(doc, "How does it integrate with existing GRC tools",
        " for automated regulatory reporting?")
    bullet(doc, "What explainability does its AI provide",
        " for regulators and auditors who will ask?")

    # SECTION 3: How AI Powers Zero-Trust Compliance in Practice
    heading(doc, "How AI Powers Zero-Trust Compliance in Practice")

    para(doc,
        "The practical value shows up in three measurable ways for compliance and security teams."
    )

    para(doc, "Real-time monitoring instead of batch audits:")
    bullet(doc, "AI continuously checks",
        " user behavior, access patterns, and data flows across the organisation")
    bullet(doc, "High-risk sessions or privilege escalations",
        " are flagged the moment they occur, not surfaced in a quarterly review")
    bullet(doc, "Evidence packs for compliance teams",
        " are auto-triggered when regulatory thresholds are breached")

    para(doc, "Adaptive, context-aware access controls:")
    bullet(doc, "Static rules don’t reflect the real world",
        " — AI adjusts access based on live risk scores, device health, and peer-group behaviour")
    bullet(doc, "Controls tighten automatically",
        " for high-risk accounts pending analyst review, without disrupting legitimate users")

    para(doc, "Human-risk scoring and insider-threat visibility:")
    bullet(doc, "Behavioural data",
        " aggregated across HR, IT, and security systems surfaces anomalous patterns")
    bullet(doc, "Unusual data exports or access outside normal hours",
        " appear as prioritised cases, not noise")
    bullet(doc, "In regulated industries",
        ", one flagged insider can trigger regulatory scrutiny — early detection changes the outcome")

    # SECTION 4: Making AI-Driven RegTech Sustainable
    heading(doc, "Making AI-Driven RegTech Sustainable at Board Level")

    para(doc,
        "The real test for the C-suite is sustainability: can your AI-driven RegTech and zero-trust "
        "compliance stack grow with the business, adapt to new regulations, and remain auditable?"
    )

    para(doc, "Three practices determine the answer:")
    bullet(doc, "Embed AI-compliance by design",
        ": Start with zero-trust principles in architecture — not retrofitted after the fact")
    bullet(doc, "Choose extensible platforms",
        ": AI-enhanced zero-trust controls must integrate with existing GRC, IAM, and SIEM tools")
    bullet(doc, "Prioritise explainability",
        ": AI models must produce reasoning regulators can understand — risk scores and access "
        "decisions need to be defensible under scrutiny")

    para(doc,
        "Over time, this approach turns AI in RegTech into a competitive advantage — cutting "
        "compliance costs, improving security posture, and building the kind of evidence that stands "
        "up to regulatory scrutiny without last-minute scrambles."
    )

    # TAKEAWAY
    heading(doc, "The Architecture Advantage for Regulated Industries")

    para(doc,
        "The firms winning on zero-trust compliance in regulated industries are not the ones with "
        "the biggest security budgets. They’re the ones that made AI and zero-trust a "
        "compliance design decision before they needed it."
    )
    para(doc,
        "This approach turns compliance from a cost centre into a competitive advantage — and "
        "builds the kind of RegTech posture that stands up to regulatory scrutiny without "
        "last-minute remediation."
    )

    # CTA
    add_horizontal_rule(doc)
    doc.add_paragraph()

    cta_p = doc.add_paragraph()
    cta_r = cta_p.add_run(
        "\U0001f517 Want the full breakdown?\n"
        "Read the complete guide on Codiste’s blog: [Insert blog URL here]\n\n"
        "\U0001f91d Modernising your RegTech stack or moving toward zero-trust compliance?\n"
        "Connect with Codiste — we help enterprises in regulated industries design, build, "
        "and scale AI-driven, zero-trust-compliant architectures that stand up to regulatory "
        "scrutiny. Reach out at codiste.com or drop a comment below."
    )
    set_font(cta_r, size=11, color=(40, 40, 40))
    doc.add_paragraph()

    # CAPTION
    section_divider(doc, "\U0001f4ac LINKEDIN CAPTION")

    cap1 = doc.add_paragraph()
    c1r = cap1.add_run(CAPTION_LINE1)
    set_font(c1r, size=12, bold=True)

    cap2 = doc.add_paragraph()
    c2r = cap2.add_run(CAPTION_LINE2)
    set_font(c2r, size=11)
    doc.add_paragraph()

    hash2 = doc.add_paragraph()
    h2r = hash2.add_run(CAPTION_HASHTAGS)
    set_font(h2r, size=11, color=(0, 100, 200))
    doc.add_paragraph()

    # PUBLISHING NOTES
    section_divider(doc, "\U0001f4cb PUBLISHING NOTES")

    notes = [
        "Word count: ~760 words",
        "Estimated read time: 3–4 min",
        "Remember to: add the blog URL in the CTA block",
        "Banner files: codiste-banner-zero-trust-compliance-ai-regtech.png / .jpg",
        "Generated: " + datetime.date.today().strftime('%d %B %Y'),
        "Target audience: CISOs, CTOs, compliance officers, C-suite in regulated industries",
    ]
    for note in notes:
        p = doc.add_paragraph()
        r = p.add_run("• " + note)
        set_font(r, size=10, color=(80, 80, 80))
        p.paragraph_format.space_after = Pt(4)

    out_path = OUT_DIR + 'LinkedIn_Article_Zero_Trust_Compliance_AI_RegTech_Codiste.docx'
    doc.save(out_path)
    print("DOCX saved: " + out_path)
    return out_path


build_docx()
