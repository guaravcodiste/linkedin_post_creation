# -*- coding: utf-8 -*-
"""
Codiste LinkedIn Article — DOCX Generator
Article: AI Decision-Making in Financial Risk Systems
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT_DIR = '/home/user/linkedin_post_creation/output/'

TITLE = "Why Your AI Risk Model Isn't the Problem: What Banks Get Wrong"

HASHTAGS_ARTICLE = "#AIInFinance #FinancialRisk #BankingAI #RegulatoryCompliance #AIGovernance"

CAPTION_LINE1 = (
    "The compliance team can't explain your AI decisions to regulators "
    "— that's not a model problem, it's an architecture problem."
)
CAPTION_LINE2 = (
    "Most financial services teams are solving the wrong thing — here's what the "
    "examination-ready ones build differently → [article link]"
)
CAPTION_HASHTAGS = "#AIRisk #BankingTech #FinancialServices #AIGovernance #FintechCompliance"


def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D7D7D7')
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_divider(doc, label):
    div1 = doc.add_paragraph()
    div1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = div1.add_run("=" * 50)
    set_font(r1, size=10, color=(100, 100, 100))

    lbl_p = doc.add_paragraph()
    lbl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = lbl_p.add_run(label)
    set_font(rl, size=11, bold=True, color=(60, 60, 60))

    div2 = doc.add_paragraph()
    div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = div2.add_run("=" * 50)
    set_font(r2, size=10, color=(100, 100, 100))

    doc.add_paragraph()


def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # ── Header ────────────────────────────────────────────────────
    section_divider(doc, "LINKEDIN ARTICLE — CODISTE")

    # ── Title ─────────────────────────────────────────────────────
    lbl = doc.add_paragraph()
    r = lbl.add_run("TITLE:")
    set_font(r, size=10, bold=True, color=(120, 120, 120))

    tp = doc.add_paragraph()
    tr = tp.add_run(TITLE)
    set_font(tr, size=20, bold=True, color=(1, 1, 1))
    doc.add_paragraph()

    lbl2 = doc.add_paragraph()
    r2 = lbl2.add_run("KEYWORDS TO USE AS HASHTAGS:")
    set_font(r2, size=10, bold=True, color=(120, 120, 120))

    hp = doc.add_paragraph()
    hr = hp.add_run(HASHTAGS_ARTICLE)
    set_font(hr, size=11, color=(0, 100, 200))
    doc.add_paragraph()

    # ── Article body ───────────────────────────────────────────────
    section_divider(doc, "ARTICLE BODY")

    # HOOK
    hook_p = doc.add_paragraph()
    hook_r = hook_p.add_run(
        "Your AI credit model scores perfectly in testing. In production, your compliance "
        "team can't explain a single decision to regulators — not because the decisions are "
        "wrong, but because the architecture around them wasn't built for examination. "
        "The model worked. The system didn't."
    )
    set_font(hook_r, size=12, bold=True)
    hook_p.paragraph_format.space_after = Pt(12)

    # SECTION 1
    def heading(doc, text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, size=14, bold=True, color=(1, 1, 60))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)

    def para(doc, text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, size=11)
        p.paragraph_format.space_after = Pt(8)

    def bullet(doc, lead, rest=""):
        p = doc.add_paragraph(style='List Bullet')
        if lead:
            rb = p.add_run(lead + (":" if rest else ""))
            set_font(rb, size=11, bold=True)
        if rest:
            rr = p.add_run(rest)
            set_font(rr, size=11)
        p.paragraph_format.space_after = Pt(5)

    heading(doc, "The Architecture Problem No One Talks About")

    para(doc,
        "McKinsey estimates AI could add $200-340 billion in annual value to global banking "
        "- roughly 9-15% of operating profits. That potential is real. But it assumes "
        "production-ready systems, not proof-of-concept models that stall before examination."
    )

    para(doc,
        "Most AI financial risk projects start with model selection. The team benchmarks "
        "performance, builds a POC, and presents strong results. Then the production deployment "
        "surfaces problems the POC was never designed to catch. The gap between a model that "
        "works and an AI risk system that passes examination is an architecture gap - and it "
        "shows up in four specific places:"
    )

    bullet(doc, "Explainability",
        " Regulators require automated credit and risk decisions to be defensible to the "
        "examiner, not just accurate in aggregate. A model that cannot produce the reasoning "
        "behind each decision is not deployable in a regulated environment."
    )
    bullet(doc, "Data lineage",
        " Every feature feeding a risk score needs a documented trail from source to feature "
        "computation - traceable on demand when an examiner asks."
    )
    bullet(doc, "Drift monitoring",
        " A model trained in one economic environment degrades as conditions change. Production "
        "systems need automated detection that flags problems before downstream risk events "
        "surface them."
    )
    bullet(doc, "Decision audit trail",
        " Every AI risk decision needs a logged record of model version, input features, and "
        "score threshold at the time - this is minimum architecture, not a reporting add-on."
    )

    para(doc,
        "None of these are post-deployment additions. They need to be designed into the system "
        "architecture before the first production inference runs."
    )

    # SECTION 2
    heading(doc, "Ethical AI in Finance Is an Engineering Problem, Not a Policy Decision")

    para(doc,
        "Bias in credit scoring, opacity in underwriting, disparate impact across demographic "
        "groups - these are well-documented challenges in AI financial risk management. They are "
        "not resolved by selecting an ethical AI framework and attaching it to an existing model. "
        "They are resolved by building specific technical controls into the development and "
        "deployment pipeline."
    )

    bullet(doc, "Bias detection",
        " Run this on training data before model training begins. Protected class proxies in "
        "zip codes, names, or spending patterns produce biased outputs even when those features "
        "are explicitly excluded."
    )
    bullet(doc, "Disparate impact testing",
        " Test for differential outcomes across demographic groups using regulatory-grade "
        "statistical methods before production deployment, with results documented for examination."
    )
    bullet(doc, "Human review pathways",
        " Build accessible recourse for any AI-assisted decision that negatively affects a "
        "consumer - the automation is appropriate; the absence of recourse is not."
    )
    bullet(doc, "Model cards",
        " Maintain version-controlled documentation of training data, performance across "
        "demographic subgroups, known limitations, and intended use cases alongside every "
        "model version."
    )

    para(doc,
        "Firms working with AI consulting partners should expect these controls in the "
        "architecture specification - not as retrospective additions if a regulatory "
        "question arises."
    )

    # SECTION 3
    heading(doc, "How Risk Domains Shape What You Actually Build")

    para(doc,
        "The specific implementation of AI decision-making varies by risk domain, but the "
        "underlying infrastructure requirements are consistent."
    )

    bullet(doc, "Credit risk",
        " The model provides the score; the decision engine applies it against a policy "
        "matrix encoding the institution's risk appetite and regulatory obligations. Keeping "
        "these layers separate means policy changes don't require model retraining, and model "
        "updates don't require policy review."
    )
    bullet(doc, "Fraud detection",
        " Decision time is measured in milliseconds. Architecture priorities shift toward "
        "inference latency, feature computation speed, and false positive rate management. "
        "The operating point on the precision-recall curve is a business decision, not a "
        "model decision."
    )
    bullet(doc, "Market risk",
        " AI tools extend what is computationally feasible in stress testing, scenario "
        "generation, and correlation analysis. Regulatory requirements for model validation "
        "are stringent - documentation needs to be part of the development process from "
        "the start."
    )

    # SECTION 4
    heading(doc, "The 3 Architecture Decisions That Define the Next 18 Months")

    para(doc,
        "The architecture decisions made now will determine whether organisations can adopt "
        "next-generation AI risk capabilities - or spend those months refactoring technical debt."
    )

    bullet(doc, "Real-time compliance integration",
        " Compliance monitoring is moving toward integration with risk scoring systems. "
        "A transaction that triggers a risk score will also trigger simultaneous checks against "
        "sanctions lists and jurisdiction-specific obligation registers - this requires risk "
        "and compliance systems to share a data infrastructure layer, a design decision "
        "that needs to be made now."
    )
    bullet(doc, "Audit-native databases",
        " The firms building AI decision workflows are choosing architectures that generate "
        "immutable decision logs, time-stamped feature vectors, and model version records "
        "as a native output - not as an afterthought."
    )
    bullet(doc, "Multimodal risk signals",
        " Voice, document, and behavioural data are entering financial risk feature sets "
        "where regulatory guidance permits. Architecture designed only for structured transaction "
        "data will require significant refactoring to extend. Building the feature engineering "
        "pipeline with extensibility in mind costs nothing now and avoids that rework later."
    )

    # TAKEAWAY
    heading(doc, "The Gap That Separates Deployable Systems From Proof of Concepts")

    para(doc,
        "The gap between a model that works and an AI risk system that passes examination is "
        "an architecture gap. Every financial services organisation moving toward AI-driven "
        "risk decisions faces it. The ones that close it before production deployment avoid "
        "the costly refactoring that comes when a model is live but not examination-ready."
    )

    para(doc,
        "The scoping conversation for a production AI risk system starts with what the "
        "current architecture is missing. Not with which model to use."
    )

    # CTA
    add_horizontal_rule(doc)
    doc.add_paragraph()

    cta_p = doc.add_paragraph()
    cta_r = cta_p.add_run(
        "Want the full breakdown?\n"
        "Read the complete guide on Codiste's blog: [Insert blog URL here]\n\n"
        "Building AI decision-making infrastructure for your financial services organisation?\n"
        "Connect with Codiste - we help banks, insurers, fintech companies, and credit unions "
        "design, build, and scale AI risk systems that meet production and examination standards. "
        "Reach out at codiste.com or drop a comment below."
    )
    set_font(cta_r, size=11, color=(40, 40, 40))
    doc.add_paragraph()

    # ── Caption ────────────────────────────────────────────────────
    section_divider(doc, "LINKEDIN CAPTION")

    cap1 = doc.add_paragraph()
    c1r = cap1.add_run(CAPTION_LINE1)
    set_font(c1r, size=12, bold=True)

    cap2 = doc.add_paragraph()
    c2r = cap2.add_run(CAPTION_LINE2)
    set_font(c2r, size=11)
    doc.add_paragraph()

    hash2 = doc.add_paragraph()
    h2r = hash2.add_run(CAPTION_HASHTAGS)
    set_font(h2r, size=11, color=(0, 100, 200))
    doc.add_paragraph()

    # ── Publishing notes ───────────────────────────────────────────
    section_divider(doc, "PUBLISHING NOTES")

    notes = [
        "Word count: ~760 words",
        "Estimated read time: 3-4 min",
        "Remember to: add the blog URL in the CTA block",
        "Banner files: codiste-banner-ai-risk-model-banks.png / .jpg",
        "Generated: " + datetime.date.today().strftime('%d %B %Y'),
        "Target audience: CTOs, fintech leaders, bank executives, compliance officers",
    ]
    for note in notes:
        p = doc.add_paragraph()
        r = p.add_run("• " + note)
        set_font(r, size=10, color=(80, 80, 80))
        p.paragraph_format.space_after = Pt(4)

    out_path = OUT_DIR + 'LinkedIn_Article_AI_Risk_Banking_Codiste.docx'
    doc.save(out_path)
    print("DOCX saved: " + out_path)
    return out_path


build_docx()
