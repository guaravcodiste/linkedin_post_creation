# -*- coding: utf-8 -*-
"""
Codiste LinkedIn Article — DOCX Generator
Row 1: Scaling Compliance Across Global Frameworks
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUT_DIR = '/home/user/linkedin_post_creation/output/scaling-compliance-global-frameworks/'
os.makedirs(OUT_DIR, exist_ok=True)

TITLE = "Scaling Compliance: Why Your Framework Breaks at Market Three"
HASHTAGS_ARTICLE = "#GlobalCompliance #RegTech #ComplianceArchitecture #SaaSScaling #RegulatoryTechnology"

def unicode_bold(text):
    result = ''
    for c in text:
        if 'A' <= c <= 'Z':
            result += chr(0x1D5D4 + ord(c) - ord('A'))
        elif 'a' <= c <= 'z':
            result += chr(0x1D5EE + ord(c) - ord('a'))
        else:
            result += c
    return result

CAPTION_LINE1 = unicode_bold(
    "Most compliance strategies are single-jurisdiction implementations wearing a framework's clothes "
    "— the difference shows up at market three."
)
CAPTION_LINE2 = (
    "Here's the three-layer architecture that makes global market entry cheaper every time, "
    "not more expensive → [article link]"
)
CAPTION_HASHTAGS = "#GlobalCompliance #ScalingCompliance #RegTech #SaaSGrowth #ComplianceStrategy"


def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D7D7D7')
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_divider(doc, label):
    div1 = doc.add_paragraph()
    div1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = div1.add_run("=" * 50)
    set_font(r1, size=10, color=(100, 100, 100))
    lbl_p = doc.add_paragraph()
    lbl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = lbl_p.add_run(label)
    set_font(rl, size=11, bold=True, color=(60, 60, 60))
    div2 = doc.add_paragraph()
    div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = div2.add_run("=" * 50)
    set_font(r2, size=10, color=(100, 100, 100))
    doc.add_paragraph()


def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=14, bold=True, color=(1, 1, 60))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=11)
    p.paragraph_format.space_after = Pt(8)


def bullet(doc, lead, rest=""):
    p = doc.add_paragraph(style='List Bullet')
    if lead:
        rb = p.add_run(lead + (":" if rest else ""))
        set_font(rb, size=11, bold=True)
    if rest:
        rr = p.add_run(rest)
        set_font(rr, size=11)
    p.paragraph_format.space_after = Pt(5)


def build_docx():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    section_divider(doc, "\U0001f537 LINKEDIN ARTICLE — CODISTE")

    lbl = doc.add_paragraph()
    r = lbl.add_run("TITLE:")
    set_font(r, size=10, bold=True, color=(120, 120, 120))

    tp = doc.add_paragraph()
    tr = tp.add_run(TITLE)
    set_font(tr, size=20, bold=True, color=(1, 1, 1))
    doc.add_paragraph()

    lbl2 = doc.add_paragraph()
    r2 = lbl2.add_run("KEYWORDS TO USE AS HASHTAGS:")
    set_font(r2, size=10, bold=True, color=(120, 120, 120))

    hp = doc.add_paragraph()
    hr = hp.add_run(HASHTAGS_ARTICLE)
    set_font(hr, size=11, color=(0, 100, 200))
    doc.add_paragraph()

    section_divider(doc, "ARTICLE BODY")

    # HOOK
    hook_p = doc.add_paragraph()
    hook_r = hook_p.add_run(
        "Most enterprises think they have a compliance strategy. What they actually have is a "
        "single-jurisdiction implementation — one that looks like a framework until market three.\n\n"
        "By the third new market, the pattern is impossible to ignore. Each new jurisdiction costs "
        "70–80% of the effort of the first. That’s not regulatory complexity. "
        "That’s an architecture failure hiding in plain sight."
    )
    set_font(hook_r, size=12, bold=True)
    hook_p.paragraph_format.space_after = Pt(12)

    # SECTION 1: Why Global Compliance Frameworks Fail
    heading(doc, "Why Global Compliance Frameworks Fail Before They Scale")

    para(doc,
        "A SaaS company closes its Series B and expands into three new markets in eighteen months. "
        "The product team hits their roadmap. The compliance team spends the same period rebuilding "
        "from scratch for each jurisdiction — at a cost that nobody modelled in the growth plan."
    )
    para(doc, "The root cause is structural:")
    bullet(doc, "The initial build",
        " didn’t separate transferable compliance logic from jurisdiction-specific requirements")
    bullet(doc, "Each new market",
        " gets treated as a ground-up rebuild, not a configuration layer on shared infrastructure")
    bullet(doc, "The architecture",
        " was built for one regulatory context — and every new context demands its own version")

    para(doc,
        "According to PwC, 85% of executives globally say compliance requirements have become more "
        "complex in the last three years, rising to 90% in financial services. The answer isn’t "
        "more compliance headcount per market. It’s architecture designed to scale before market two."
    )

    # SECTION 2: The Three-Layer Architecture
    heading(doc, "The Three-Layer Architecture That Actually Scales")

    para(doc,
        "Organisations that enter new markets without proportional compliance cost growth share one "
        "structural characteristic: they separated three layers explicitly before entering their "
        "second market."
    )

    bullet(doc, "Core compliance infrastructure",
        ": Audit trail generation, policy versioning, and reporting pipelines — built once, "
        "deployed across all markets without modification")
    bullet(doc, "Jurisdictional configuration",
        ": Data residency rules, consent formats, and reporting deadlines — configured per "
        "market as an additive layer on top of shared infrastructure")
    bullet(doc, "Regulatory monitoring",
        ": Change detection, obligation mapping, and threshold alerts — market-specific in "
        "its inputs, shared in its architecture")

    para(doc,
        "This separation is an architecture decision, not a platform purchase. It needs to be made "
        "before selecting a compliance platform — not after the scaling problem surfaces at "
        "market three. Global compliance architecture built for scale treats each new market as a "
        "configuration layer, not a rebuild."
    )

    # SECTION 3: What Scalable Compliance Services Need
    heading(doc, "What Scalable Compliance Services Actually Need to Deliver")

    para(doc,
        "Most compliance platforms solve the first-market problem well. The scaling problem, they "
        "solve poorly. When a client enters a new jurisdiction, the platform either lacks coverage "
        "or requires a vendor engagement to extend it."
    )
    para(doc, "The features that separate scalable from non-scalable compliance architecture:")

    bullet(doc, "Jurisdictional rule engine",
        ": Regulatory requirements expressed as configurable rules, not hard-coded logic — "
        "new markets require rule additions, not code changes")
    bullet(doc, "Policy versioning with rollback",
        ": Regulatory guidance changes should be a content management operation, not a system deployment")
    bullet(doc, "Consent and data residency routing",
        ": Data routed to the correct processing environment at point of capture, not retroactively at reporting")
    bullet(doc, "Multi-jurisdiction reporting pipeline",
        ": One abstracted submission layer so compliance teams manage content, not infrastructure")

    # SECTION 4: Scaling Without Growing Headcount
    heading(doc, "How Leading Teams Scale Without Growing Compliance Headcount")

    para(doc,
        "The compliance teams managing global regulatory obligations without proportional headcount "
        "growth have made one critical distinction: they separate the work that requires human "
        "judgment from the work that doesn’t."
    )
    para(doc,
        "Human judgment is required for regulatory interpretation, responding to regulatory "
        "inquiries, and assessing whether a flagged risk pattern warrants escalation. It is not "
        "required for evidence collection, report generation, audit trail assembly, policy "
        "distribution, or consent capture. These should be automated as compliance infrastructure."
    )
    para(doc,
        "Gartner projects that fragmented AI regulation will cover 50% of the world’s "
        "economies by 2027, driving $5 billion in compliance investment — disproportionately "
        "hitting firms without scalable architecture."
    )
    para(doc, "The organisations winning at global compliance have built systems where:")
    bullet(doc, "Regulatory change monitoring",
        " feeds directly into obligation mapping, reducing manual tracking")
    bullet(doc, "Evidence",
        " is collected automatically from source systems, not assembled manually before each audit")
    bullet(doc, "Compliance dashboards",
        " surface the current status of every regulatory obligation by jurisdiction")

    # TAKEAWAY
    heading(doc, "The Architecture Builds Before Market Three or Pays for It After")

    para(doc,
        "Every market entered without a scalable compliance architecture is a market that will need "
        "retrofitting later — at a higher cost. The firms growing fastest across multiple "
        "jurisdictions didn’t hire more compliance headcount per market. They built "
        "infrastructure that made each entry cheaper than the last."
    )
    para(doc,
        "That architecture is buildable on what most organisations already have. The gap is usually "
        "not technology — it’s the design decision separating the transferable from the "
        "jurisdiction-specific, made before market two forces the issue."
    )

    # CTA
    add_horizontal_rule(doc)
    doc.add_paragraph()

    cta_p = doc.add_paragraph()
    cta_r = cta_p.add_run(
        "\U0001f517 Want the full breakdown?\n"
        "Read the complete guide on Codiste’s blog: [Insert blog URL here]\n\n"
        "\U0001f91d Scaling compliance into new markets?\n"
        "Connect with Codiste — we help SaaS companies, e-commerce platforms, and multinational "
        "enterprises design, build, and scale compliance architectures that make each market entry "
        "cheaper than the last. Reach out at codiste.com or drop a comment below."
    )
    set_font(cta_r, size=11, color=(40, 40, 40))
    doc.add_paragraph()

    # CAPTION
    section_divider(doc, "\U0001f4ac LINKEDIN CAPTION")

    cap1 = doc.add_paragraph()
    c1r = cap1.add_run(CAPTION_LINE1)
    set_font(c1r, size=12, bold=True)

    cap2 = doc.add_paragraph()
    c2r = cap2.add_run(CAPTION_LINE2)
    set_font(c2r, size=11)
    doc.add_paragraph()

    hash2 = doc.add_paragraph()
    h2r = hash2.add_run(CAPTION_HASHTAGS)
    set_font(h2r, size=11, color=(0, 100, 200))
    doc.add_paragraph()

    # PUBLISHING NOTES
    section_divider(doc, "\U0001f4cb PUBLISHING NOTES")

    notes = [
        "Word count: ~750 words",
        "Estimated read time: 3–4 min",
        "Remember to: add the blog URL in the CTA block",
        "Banner files: codiste-banner-scaling-compliance-global-frameworks.png / .jpg",
        "Generated: " + datetime.date.today().strftime('%d %B %Y'),
        "Target audience: CTOs, SaaS founders, compliance officers, multinational enterprise leaders",
    ]
    for note in notes:
        p = doc.add_paragraph()
        r = p.add_run("• " + note)
        set_font(r, size=10, color=(80, 80, 80))
        p.paragraph_format.space_after = Pt(4)

    out_path = OUT_DIR + 'LinkedIn_Article_Scaling_Compliance_Global_Frameworks_Codiste.docx'
    doc.save(out_path)
    print("DOCX saved: " + out_path)
    return out_path


build_docx()
