# -*- coding: utf-8 -*-
"""
Codiste LinkedIn Article — DOCX Generator
Topic: AI Compliance Ethics & Governance
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUT_DIR = '/home/user/linkedin_post_creation/output/ai-compliance-ethics-governance/'
os.makedirs(OUT_DIR, exist_ok=True)

TITLE = "AI Compliance Is Scaling Fast: Your Ethics Framework Isn't"

HASHTAGS_ARTICLE = "#AICompliance #EthicalAI #AIGovernance #RegulatoryCompliance #AIPrivacy"

CAPTION_LINE1 = (
    "\U0001d5f6\U0001d5f3\U0001d5f9\U0001d606 35% of companies have an AI governance framework "
    "— yet AI already makes compliance decisions regulators will ask you to explain."
)
CAPTION_LINE2 = (
    "Ethics alone won’t close the gap — here’s the governance architecture "
    "that actually does → [article link]"
)
CAPTION_HASHTAGS = "#AICompliance #EthicalAI #AIGovernance #GDPR #RegulatoryTech"


def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D7D7D7')
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_divider(doc, label):
    div1 = doc.add_paragraph()
    div1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = div1.add_run("=" * 50)
    set_font(r1, size=10, color=(100, 100, 100))

    lbl_p = doc.add_paragraph()
    lbl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = lbl_p.add_run(label)
    set_font(rl, size=11, bold=True, color=(60, 60, 60))

    div2 = doc.add_paragraph()
    div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = div2.add_run("=" * 50)
    set_font(r2, size=10, color=(100, 100, 100))

    doc.add_paragraph()


def heading(doc, draw, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=14, bold=True, color=(1, 1, 60))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=11)
    p.paragraph_format.space_after = Pt(8)


def bullet(doc, lead, rest=""):
    p = doc.add_paragraph(style='List Bullet')
    if lead:
        rb = p.add_run(lead + (":" if rest else ""))
        set_font(rb, size=11, bold=True)
    if rest:
        rr = p.add_run(rest)
        set_font(rr, size=11)
    p.paragraph_format.space_after = Pt(5)


def build_docx():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    section_divider(doc, "LINKEDIN ARTICLE — CODISTE")

    lbl = doc.add_paragraph()
    r = lbl.add_run("TITLE:")
    set_font(r, size=10, bold=True, color=(120, 120, 120))

    tp = doc.add_paragraph()
    tr = tp.add_run(TITLE)
    set_font(tr, size=20, bold=True, color=(1, 1, 1))
    doc.add_paragraph()

    lbl2 = doc.add_paragraph()
    r2 = lbl2.add_run("KEYWORDS TO USE AS HASHTAGS:")
    set_font(r2, size=10, bold=True, color=(120, 120, 120))

    hp = doc.add_paragraph()
    hr = hp.add_run(HASHTAGS_ARTICLE)
    set_font(hr, size=11, color=(0, 100, 200))
    doc.add_paragraph()

    section_divider(doc, "ARTICLE BODY")

    # HOOK
    hook_p = doc.add_paragraph()
    hook_r = hook_p.add_run(
        "AI is quietly reviewing contracts, screening hires, and flagging suspicious "
        "transactions at thousands of companies right now. The efficiency gains are real. "
        "But most of those deployments are running without the ethical architecture to make "
        "their decisions defensible — to regulators, to the people they affect, and to "
        "any auditor who asks why."
    )
    set_font(hook_r, size=12, bold=True)
    hook_p.paragraph_format.space_after = Pt(12)

    # CONTEXT
    draw = None  # placeholder for heading calls

    def h(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, size=14, bold=True, color=(1, 1, 60))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)

    para(doc,
        "Only 35% of companies currently have an AI governance framework in place, and "
        "fewer than 20% conduct regular AI audits. That gap is where the real risk lives. "
        "Unlike rule-based systems, AI in compliance learns from data — which means it "
        "can reproduce bias, obscure its reasoning, and handle sensitive personal information "
        "in ways that create serious legal exposure before anyone notices."
    )

    para(doc, "The risks compound across three dimensions:")

    bullet(doc, "Reputational",
        " An AI-powered hiring tool that discriminates can trigger public backlash and "
        "regulatory scrutiny within hours of exposure."
    )
    bullet(doc, "Financial and regulatory",
        " GDPR, CCPA, and emerging AI-specific laws are rapidly tightening what ‘compliance’ "
        "actually requires from automated decision systems."
    )
    bullet(doc, "Operational",
        " If compliance teams don’t trust the AI model, they ignore alerts or override outputs "
        "— breaking the very workflow the AI was built to strengthen."
    )

    h("Three Ethical Risks Compliance Teams Underestimate")

    para(doc,
        "Bias and fairness failures don’t announce themselves. Many AI compliance tools are "
        "trained on historical data carrying systemic biases around gender, race, geography, or "
        "income. An anti-money-laundering engine that disproportionately flags customers from "
        "certain regions doesn’t just fail ethically — it may violate anti-discrimination rules."
    )

    para(doc, "Common patterns to watch for:")
    bullet(doc, "Uneven risk-scoring", " across customer or employee demographic groups")
    bullet(doc, "Hidden proxy variables",
        " — zip code acting as a proxy for race, for example"
    )
    bullet(doc, "No measurable fairness metrics", " tracked or audited over time")

    para(doc,
        "Transparency is the second gap. If a compliance AI flags a transaction or employee "
        "behavior, regulators and affected individuals expect to understand why. A black-box "
        "model with no human-readable justification cannot satisfy an auditor and cannot be "
        "defended when challenged."
    )

    para(doc,
        "Privacy is the third. AI compliance systems rely heavily on personal data. Without "
        "privacy-by-design — data minimization, anonymization, and clear legal bases for "
        "processing — every workflow carries an ethics liability embedded in its architecture."
    )

    h("What a Real Ethical AI Policy Looks Like")

    para(doc,
        "86% of users prefer brands with transparent AI policies. That’s not just a "
        "compliance obligation — it’s a direct driver of customer trust. Yet most "
        "organizations treat “ethical AI” as a public positioning statement rather than "
        "an operational specification with teeth."
    )

    para(doc, "A policy that actually functions includes five pillars:")
    bullet(doc, "Fairness and non-discrimination clauses",
        " — how models are tested, how fairness is measured, and who is accountable "
        "for disparate outcomes"
    )
    bullet(doc, "Transparency obligations",
        " — what level of explanation is required for each decision type, and to whom "
        "(regulators, customers, HR, auditors)"
    )
    bullet(doc, "Data protection commitments",
        " — aligned with GDPR and CCPA, including data-retention limits and opt-out rights"
    )
    bullet(doc, "Human oversight and escalation paths",
        " — when AI outputs must be reviewed, challenged, or overridden, and how "
        "decisions are logged"
    )
    bullet(doc, "Model lifecycle governance",
        " — from training through retirement, with assigned roles, review cycles, "
        "and rollback mechanisms"
    )

    h("AI Governance: Turning Ethics Into Daily Operations")

    para(doc,
        "AI governance is the bridge between high-level ethics principles and day-to-day "
        "operations. It answers the questions auditors actually ask: Who owns the model? "
        "Who approved it? How was it monitored? What happened when it drifted?"
    )

    para(doc, "Getting governance off whiteboards and into workflows requires:")
    bullet(doc, "A cross-functional AI governance committee",
        " — compliance, legal, risk, data science, and IT all need representation"
    )
    bullet(doc, "Risk-tier classification for AI use cases",
        " — hiring, credit scoring, and fraud detection warrant stricter review cycles "
        "than lower-risk document summarization"
    )
    bullet(doc, "Mapped workflows with escalation rules",
        " and full audit trails at every decision handoff"
    )
    bullet(doc, "Production monitoring dashboards",
        " tracking model drift, fairness metrics, and incident rates over time"
    )
    bullet(doc, "AI literacy training for non-technical staff",
        " — HR teams, controllers, and compliance officers need enough understanding "
        "to challenge AI-driven decisions when something looks wrong"
    )

    # TAKEAWAY
    h("The Gap Between Capability and Accountability")

    para(doc,
        "AI in compliance offers measurable benefits when it’s backed by genuine ethics "
        "and governance infrastructure. The organizations getting it right aren’t just "
        "deploying more capable AI — they’re building AI that survives regulatory "
        "scrutiny, earns employee trust, and holds up when someone asks hard questions."
    )

    para(doc,
        "The question for every enterprise scaling AI in compliance isn’t whether the "
        "model performs. It’s whether your governance structure can defend every "
        "decision it makes."
    )

    add_horizontal_rule(doc)
    doc.add_paragraph()

    cta_p = doc.add_paragraph()
    cta_r = cta_p.add_run(
        "\U0001f517 Want the full breakdown?\n"
        "Read the complete guide on Codiste’s blog: [Insert blog URL here]\n\n"
        "\U0001f91d Building or upgrading AI compliance infrastructure?\n"
        "Connect with Codiste — we help global enterprises design, build, and scale "
        "AI solutions aligned with GDPR, CCPA, and emerging AI-specific regulatory standards. "
        "Reach out at codiste.com or drop a comment below."
    )
    set_font(cta_r, size=11, color=(40, 40, 40))
    doc.add_paragraph()

    section_divider(doc, "LINKEDIN CAPTION")

    cap1 = doc.add_paragraph()
    c1r = cap1.add_run(CAPTION_LINE1)
    set_font(c1r, size=12, bold=True)

    cap2 = doc.add_paragraph()
    c2r = cap2.add_run(CAPTION_LINE2)
    set_font(c2r, size=11)
    doc.add_paragraph()

    hash2 = doc.add_paragraph()
    h2r = hash2.add_run(CAPTION_HASHTAGS)
    set_font(h2r, size=11, color=(0, 100, 200))
    doc.add_paragraph()

    section_divider(doc, "PUBLISHING NOTES")

    notes = [
        "Word count: ~760 words",
        "Estimated read time: 3-4 min",
        "Remember to: add the blog URL in the CTA block",
        "Banner files: codiste-banner-ai-compliance-ethics-governance.png / .jpg",
        "Generated: " + datetime.date.today().strftime('%d %B %Y'),
        "Target audience: CTOs, startup founders, compliance officers, enterprise tech leaders",
    ]
    for note in notes:
        p = doc.add_paragraph()
        r = p.add_run("• " + note)
        set_font(r, size=10, color=(80, 80, 80))
        p.paragraph_format.space_after = Pt(4)

    out_path = OUT_DIR + 'AI_Compliance_Ethics_Governance_Codiste.docx'
    doc.save(out_path)
    print("DOCX saved: " + out_path)
    return out_path


build_docx()
